"""
GET /api/download?token=xxx
從 Redis 讀取已存的行程計畫，回傳 .docx Word 檔供下載
"""

import io
import json
import os
import sys
import traceback
from http.server import BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from bot.services.redis_store import redis_get


# ── 主題色系 ───────────────────────────────────────────
_THEME = {
    "primary":   "1A237E",   # 深藍
    "accent":    "0D47A1",   # 中藍
    "light":     "E3F2FD",   # 淺藍背景
    "header_bg": "1565C0",   # 表格 header 深藍
    "row_alt":   "F5F9FF",   # 表格交替行
    "label_bg":  "E8EAF6",   # 標籤底色
    "section":   "1976D2",   # Section 標題
    "day_bg":    "E1F5FE",   # Day 標題底色
    "insider":   "E8F5E9",   # 在地眉角底色（淡綠）
    "gray":      "757575",
    "divider":   "BBDEFB",
}


def _hex(name: str) -> RGBColor:
    h = _THEME[name]
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color: str = "BBDEFB"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ("w:top", "w:left", "w:bottom", "w:right"):
        border = OxmlElement(side)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        tcBorders.append(border)


def _add_section_heading(doc, text: str, bg_color: str = "1565C0"):
    """色塊背景的 section 標題"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, bg_color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_sub_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _hex("section")
    return p


def _add_info_table(doc, rows_data: list[tuple[str, str]]):
    """兩欄資訊表格（label | value）"""
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(3.5)
    tbl.columns[1].width = Cm(13)
    for i, (label, value) in enumerate(rows_data):
        row = tbl.add_row()
        lbl, val = row.cells[0], row.cells[1]
        bg = _THEME["label_bg"] if i % 2 == 0 else _THEME["row_alt"]
        _set_cell_bg(lbl, _THEME["label_bg"])
        _set_cell_bg(val, bg)
        lp = lbl.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = _hex("accent")
        vp = val.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(9)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _add_day_block(doc, day: dict):
    """每日行程 — 帶底色的日期標題 + 三格時段"""
    title = day.get("title", "")
    date_label = day.get("date_label", "")
    header = f"{title}  {date_label}".strip()

    # 日期標題列
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, _THEME["day_bg"])
    p = cell.paragraphs[0]
    run = p.add_run(f"  {header}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _hex("accent")

    # 三時段
    slots = [
        ("🌅 上午", day.get("am", "")),
        ("☀️ 下午", day.get("pm", "")),
        ("🌙 晚上", day.get("eve", "")),
    ]
    for icon, val in slots:
        if not val:
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        icon_run = p.add_run(f"{icon}  ")
        icon_run.bold = True
        icon_run.font.size = Pt(9)
        icon_run.font.color.rgb = _hex("gray")
        val_run = p.add_run(val)
        val_run.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── 機場交通資料 ────────────────────────────────────────
_AIRPORT_TRANSIT: dict[str, dict] = {
    "GMP": {
        "name": "金浦機場（GMP）",
        "to_hotel": [
            "出境大廳找「地下鐵 5 號線」入口（地下 1 樓）",
            "刷 T-money 卡進閘口，搭 5 號線（紫色）",
            "孔德站（공덕역）轉 6 號線（棕色）",
            "忠武路站（충무로역）下車，3 號出口步行 8 分鐘到明洞",
            "全程約 50 分鐘 ｜ 費用約 1,500 韓元（NT$35）",
        ],
        "tip": "不想轉車？叫 Kakao T 計程車直達飯店，約 25,000–35,000 韓元（NT$600–850）",
    },
    "ICN": {
        "name": "仁川機場（ICN）",
        "to_hotel": [
            "入境大廳 B1 搭機場鐵路（AREX）",
            "直達列車到首爾站約 43 分鐘，1,500 韓元",
            "一般列車各站停，約 66 分鐘，900 韓元",
            "首爾站轉 1 號線或 4 號線到明洞",
        ],
        "tip": "機場巴士（리무진버스）也可直達明洞，約 70 分鐘，16,000 韓元",
    },
    "NRT": {
        "name": "成田機場（NRT）",
        "to_hotel": [
            "成田特快 N'EX 到新宿/渋谷約 90 分鐘，3,070 日元",
            "京成 Skyliner 到上野/日暮里約 41 分鐘，2,570 日元",
            "利木津巴士到主要飯店區，約 60–90 分鐘",
        ],
        "tip": "建議台灣出發前購買 Suica/Pasmo IC 卡，機場自動售票機也可購買",
    },
    "HND": {
        "name": "羽田機場（HND）",
        "to_hotel": [
            "東京單軌電車到浜松町，轉 JR 山手線，約 30 分鐘",
            "京急線直通品川站約 13 分鐘，600 日元",
            "利木津巴士到主要飯店區約 30–60 分鐘",
        ],
        "tip": "計程車到新宿約 4,500–6,000 日元，距市區較近",
    },
}

_KOREA_APPS: list[tuple[str, str]] = [
    ("Kakao Map",   "韓國最準確的地圖，支援步行/地鐵/計程車路線"),
    ("Kakao T",     "叫計程車必備，可用信用卡付款"),
    ("Naver Map",   "景點評論豐富，適合查餐廳開放時間"),
    ("Seoul Metro", "首爾地鐵路線圖，可離線使用"),
    ("Papago",      "韓語翻譯、相機即時翻字幕"),
]

_JAPAN_APPS: list[tuple[str, str]] = [
    ("Google Maps",    "東京地鐵路線最準確"),
    ("Yahoo!乗換案内", "電車換乘查詢最詳細"),
    ("Tabelog",        "日本餐廳評論第一名"),
    ("じゃらん",        "景點預訂、查票價/開放時間"),
]


def _add_airport_guide_section(doc: "Document", dest_code: str) -> None:
    """✈️ 機場攻略 section：出發流程＋入境＋交通到飯店＋APP"""
    dest_upper = dest_code.upper() if dest_code else ""
    transit = _AIRPORT_TRANSIT.get(dest_upper)

    _add_section_heading(doc, "✈️  機場攻略", bg_color="37474F")

    _add_sub_heading(doc, "🛫 出發流程（台灣機場）")
    for step in [
        "出發前 2.5 小時抵達機場，看電子看板確認航班與航空公司報到櫃台",
        "Check-in：護照＋電子機票給地勤，托運行李（提前在家秤重，超重費很貴！）",
        "安全檢查：電腦、外套、皮帶放 X 光籃；液體/凝膠 ≤100ml 裝透明夾鏈袋",
        "出境查驗：走「本國人」通道遞護照，或使用 e-Gate 自動通關（須事先向移民署申辦）",
        "候機：提前 30 分鐘到登機門等候，廣播叫號後再排隊",
        "行動電源限 100Wh（≈27,000mAh）以下才能帶上機",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.add_run(step).font.size = Pt(9)

    _add_sub_heading(doc, "🛬 抵達目的地入境")
    for step in [
        "跟著「入境 / Arrivals」指標走，外國人通道（Foreigner）排隊",
        "護照遞給海關，問「旅遊目的」回答「Tourism / 觀光」",
        "看入境大廳電子看板，找航班號碼對應的行李轉盤取行李",
        "沒超額物品走綠色通道「Nothing to Declare」直接出關",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.add_run(step).font.size = Pt(9)

    if transit:
        _add_sub_heading(doc, f"🚇 {transit['name']} → 飯店")
        for step in transit["to_hotel"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.add_run(step).font.size = Pt(9)
        tip_p = doc.add_paragraph()
        tip_p.paragraph_format.left_indent = Cm(0.5)
        tip_r = tip_p.add_run(f"💡 {transit['tip']}")
        tip_r.font.size = Pt(9)
        tip_r.font.color.rgb = _hex("gray")
        tip_r.italic = True

    if dest_upper in ("GMP", "ICN"):
        _add_sub_heading(doc, "💳 T-money 交通卡")
        _add_info_table(doc, [
            ("在哪買",   "便利商店（GS25、CU、7-Eleven）或地鐵站自動售票機"),
            ("充值方式", "便利商店現金儲值，最少 1,000 韓元起"),
            ("優惠",     "刷卡比現金便宜 100 韓元，轉乘 30 分鐘內免費"),
            ("退卡",     "回台前到地鐵客服退押金 500 韓元＋剩餘餘額"),
        ])

    app_list = (
        _KOREA_APPS if dest_upper in ("GMP", "ICN") else
        _JAPAN_APPS if dest_upper in ("NRT", "HND") else []
    )
    if app_list:
        country = "韓國" if dest_upper in ("GMP", "ICN") else "日本"
        _add_sub_heading(doc, f"📱 {country}必備 APP")
        _add_info_table(doc, list(app_list))

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _build_docx(plan: dict) -> bytes:
    doc = Document()

    # ── 頁面設定 ───────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    flag      = plan.get("flag", "✈️")
    city      = plan.get("city", "")
    days_text = plan.get("days_text", "")

    # ── 主標題（深藍色塊）────────────────────────────────
    tbl_title = doc.add_table(rows=2, cols=1)
    tbl_title.style = "Table Grid"
    # Row 1：城市名
    c1 = tbl_title.rows[0].cells[0]
    _set_cell_bg(c1, _THEME["primary"])
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"{flag}  {city}  行程計畫書")
    r1.bold = True
    r1.font.size = Pt(22)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(4)
    # Row 2：日期副標
    c2 = tbl_title.rows[1].cells[0]
    _set_cell_bg(c2, _THEME["accent"])
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(days_text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xE3, 0xF2, 0xFD)
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(6)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

    # ── 基本資料 ─────────────────────────────────────────
    _add_section_heading(doc, "📋  基本資料")
    rows_base = [
        ("目的地", f"{flag}  {city}"),
        ("出發地", f"{plan.get('origin_name', '')}  →  {city}"),
        ("日期",   plan.get("date_display", "")),
        ("人數",   f"{plan.get('adults', 1)} 人"),
    ]
    if plan.get("budget"):
        rows_base.append(("預算", f"NT${plan['budget']:,}"))
    _add_info_table(doc, rows_base)

    # ── 旅遊資訊 ─────────────────────────────────────────
    info_rows = []
    if plan.get("flight_text"):
        info_rows.append(("✈️ 機票",   plan["flight_text"]))
    if plan.get("hotel_pref"):
        info_rows.append(("🏨 住宿",   plan["hotel_pref"]))
    if plan.get("visa_text"):
        info_rows.append(("📘 簽證",   plan["visa_text"]))
    if plan.get("weather_text"):
        info_rows.append(("🌤 天氣",   plan["weather_text"]))
    if plan.get("exchange_text"):
        info_rows.append(("💱 匯率",   plan["exchange_text"]))
    if plan.get("plug_text"):
        info_rows.append(("🔌 插座",   plan["plug_text"]))
    if plan.get("custom"):
        info_rows.append(("📝 特別需求", plan["custom"]))

    if info_rows:
        _add_section_heading(doc, "🗺  旅遊資訊")
        _add_info_table(doc, info_rows)

    # ── 機場攻略（有 dest_code 就加） ────────────────────
    dest_code = plan.get("dest_code", "")
    if dest_code:
        _add_airport_guide_section(doc, dest_code)

    # ── 天天行程 ─────────────────────────────────────────
    llm_itinerary = plan.get("llm_itinerary") or []
    itinerary = plan.get("itinerary", [])

    if llm_itinerary:
        _add_section_heading(doc, "📅  天天行程  ✨ AI 個人化")
        depart = plan.get("depart_date", "")
        import datetime as _dt
        for i, day in enumerate(llm_itinerary):
            date_label = ""
            if depart:
                try:
                    d = _dt.date.fromisoformat(depart[:10]) + _dt.timedelta(days=i)
                    date_label = d.strftime("%m/%d")
                except Exception:
                    pass
            _add_day_block(doc, {
                "title": f"Day {i + 1}  {day.get('theme', '')}",
                "date_label": date_label,
                "am":  day.get("am", ""),
                "pm":  day.get("pm", ""),
                "eve": day.get("eve", ""),
            })
    elif itinerary:
        _add_section_heading(doc, "📅  天天行程")
        for day in itinerary:
            _add_day_block(doc, day)

    # ── 在地眉角 ─────────────────────────────────────────
    insider = plan.get("insider", {})
    if insider:
        _add_section_heading(doc, f"✨  {city} 在地眉角", bg_color="1B5E20")
        sections = [
            ("🎟️ 票務時機",  insider.get("ticket", [])),
            ("👥 人潮規律",  insider.get("crowd", [])),
            ("🚇 交通秘技",  insider.get("transport", [])),
            ("🗺️ 隱藏景點",  insider.get("hidden", [])),
            ("💰 省錢技巧",  insider.get("money", [])),
        ]
        for sec_title, items in sections:
            if not items:
                continue
            _add_sub_heading(doc, sec_title)
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(item)
                run.font.size = Pt(9)
        sp2 = doc.add_paragraph()
        sp2.paragraph_format.space_after = Pt(4)

    # ── 必吃清單 ─────────────────────────────────────────
    must_eat = plan.get("must_eat", [])
    if must_eat:
        _add_section_heading(doc, f"🍜  {city} 必吃清單", bg_color="B71C1C")
        for item in must_eat:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(1)
            run = p.add_run(item)
            run.font.size = Pt(10)
        doc.add_paragraph()

    # ── 頁尾 ─────────────────────────────────────────────
    tbl_footer = doc.add_table(rows=1, cols=1)
    tbl_footer.style = "Table Grid"
    cf = tbl_footer.rows[0].cells[0]
    _set_cell_bg(cf, _THEME["light"])
    pf = cf.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("⚠️ 簽證/海關資訊僅供參考，以官方公告為準    |    由「出國優轉」LINE Bot 生成")
    rf.font.size = Pt(8)
    rf.font.color.rgb = _hex("gray")
    pf.paragraph_format.space_before = Pt(4)
    pf.paragraph_format.space_after = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class handler(BaseHTTPRequestHandler):
    def do_GET(self):
        try:
            self._handle_get()
        except Exception as e:
            print(f"[download] unhandled: {traceback.format_exc()}")
            try:
                self.send_response(500)
                self.end_headers()
                self.wfile.write(str(e).encode("utf-8"))
            except Exception:
                pass

    def _handle_get(self):
        parsed = urlparse(self.path)
        params = parse_qs(parsed.query)
        token  = params.get("token", [None])[0]

        if not token:
            self.send_response(400)
            self.end_headers()
            self.wfile.write(b"Missing token")
            return

        raw = redis_get(f"download:{token}")
        if not raw:
            self.send_response(404)
            self.end_headers()
            self.wfile.write("找不到行程，連結已過期（72小時）".encode("utf-8"))
            return

        try:
            plan = json.loads(raw)
        except Exception:
            self.send_response(500)
            self.end_headers()
            return

        try:
            content = _build_docx(plan)
        except Exception as e:
            self.send_response(500)
            self.end_headers()
            self.wfile.write(str(e).encode("utf-8"))
            return

        city         = plan.get("city", "行程")
        date_display = plan.get("date_display", "").replace("/", "-").replace(" ~ ", "_")
        filename     = f"{city}_{date_display}_計畫書.docx"
        from urllib.parse import quote
        encoded_name = quote(filename, safe="")

        self.send_response(200)
        self.send_header("Content-Type",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.send_header("Content-Disposition",
            f'attachment; filename="{filename}"; filename*=UTF-8\'\'{encoded_name}')
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)

    def log_message(self, *args):
        pass
